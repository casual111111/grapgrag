#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将知识图谱统计表格插入到 Word 文档中
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def add_table_to_doc(doc, title, headers, rows):
    """添加带标题的表格到文档"""
    # 添加标题（使用普通段落加粗）
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 创建表格（使用默认样式）
    table = doc.add_table(rows=1, cols=len(headers))

    # 添加表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 设置表头加粗
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # 添加数据行
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
            # 设置单元格字体大小
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # 添加空行
    doc.add_paragraph()

def add_section_title(doc, title, level=1):
    """添加章节标题"""
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    if level == 1:
        heading_run.font.size = Pt(16)
    else:
        heading_run.font.size = Pt(14)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text):
    """添加正文段落"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(11)
    return para

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        para = doc.add_paragraph()
        run = para.add_run(f"• {item}")
        run.font.size = Pt(11)

def add_neo4j_reasons(doc):
    """添加"为什么要在 Neo4j 部署"章节"""
    add_section_title(doc, '一、为什么要在 Neo4j 部署')

    add_paragraph(doc, 'GraphRAG 原来的 local search 实现存在以下局限性：')

    add_bullet_list(doc, [
        '知识图谱以 parquet 表格形式存储，不是真正的图数据库',
        '每次查询都需要线性扫描 relationships.parquet 文件查找关系，效率低下',
        '只能做 1-hop 邻居查找，无法进行多跳图遍历',
        '无法利用图算法（如 PageRank、社区发现、最短路径等）',
        '随着实体和关系数量增长，查询性能线性下降',
    ])

    add_paragraph(doc, '部署 Neo4j 的优势：')

    add_bullet_list(doc, [
        '原生图数据库存储，关系查找效率 O(1) 而非 O(N) 全表扫描',
        '通过 Cypher 查询语言支持任意跳数的图遍历（MATCH path = (a)-[*1..N]-(b)）',
        '内置向量索引（Vector Index），支持语义相似度检索',
        '可扩展性强，支持大规模知识图谱',
        '可视化图结构，便于调试和分析',
    ])

    doc.add_paragraph()  # 空行

def add_graphrag_principle(doc):
    """添加"GraphRAG 原理"章节"""
    add_section_title(doc, '二、GraphRAG 原理')

    add_section_title(doc, '2.1 索引阶段（Indexing）', level=2)
    add_paragraph(doc, 'GraphRAG 首先对原始文档进行索引，构建知识图谱：')

    add_bullet_list(doc, [
        '文档分块：将原始文档切分为 1200 token 左右的文本块（text units）',
        '实体抽取：使用 LLM 从每个文本块中抽取实体（Entity），包括概念、状态、事件、规则等',
        '关系抽取：识别实体之间的关系，包括关系描述和权重',
        '社区检测：使用 Leiden 算法对图进行社区划分，发现实体聚类',
        '社区报告：为每个社区生成总结性报告，描述该主题的核心信息',
        '向量编码：对实体描述和文本块进行 embedding，存入向量索引',
    ])

    add_section_title(doc, '2.2 查询阶段（Local Search）', level=2)
    add_paragraph(doc, '当用户提问时，GraphRAG 的 local search 执行以下流程：')

    add_bullet_list(doc, [
        '1. 向量检索：将用户问题编码为向量，在实体描述向量库中搜索最相似的 top-10 个实体',
        '2. 关系查找：线性扫描 relationships.parquet，找到 source 或 target 在这 10 个实体中的关系',
        '3. 文本块关联：从实体和关系的 text_unit_ids 字段收集相关文本块',
        '4. 上下文组装：按 token 预算比例（实体35% + 文本块65%）拼接上下文',
        '5. 生成答案：将上下文和用户问题发送给 LLM，生成最终回答',
    ])

    add_paragraph(doc, '核心局限：步骤2中的关系查找是 1-hop 的，代码实现如下：')

    # 添加代码示例
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        'entity_relationships = [\n'
        '    rel for rel in all_relationships\n'
        '    if rel.source in seed_entities or rel.target in seed_entities\n'
        ']'
    )
    code_run.font.size = Pt(9)
    code_run.font.name = 'Consolas'

    add_paragraph(doc, '这种方式只能找到直接相连的关系，无法发现"邻居的邻居"，导致检索覆盖面有限。')

    doc.add_paragraph()  # 空行

def main():
    docx_path = Path(r"C:\保存\周报\7.24.docx")

    # 检查文件是否存在
    if docx_path.exists():
        doc = Document(str(docx_path))
        print(f"打开现有文档: {docx_path}")
    else:
        doc = Document()
        print(f"创建新文档: {docx_path}")

    # 添加主标题
    main_title = doc.add_paragraph()
    main_title_run = main_title.add_run('Neo4j 知识图谱部署与 GraphRAG 原理')
    main_title_run.bold = True
    main_title_run.font.size = Pt(18)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # 空行

    # 添加"为什么要在 Neo4j 部署"章节
    add_neo4j_reasons(doc)

    # 添加"GraphRAG 原理"章节
    add_graphrag_principle(doc)

    # 添加数据统计章节
    add_section_title(doc, '三、知识图谱数据统计')
    doc.add_paragraph()

    # 节点统计表
    node_headers = ['节点类型', '数量']
    node_rows = [
        ['Entity（实体）', '4,111'],
        ['TextUnit（文本块）', '179'],
        ['Document（文档）', '14']
    ]
    add_table_to_doc(doc, '3.1 节点统计', node_headers, node_rows)

    # 关系统计表
    rel_headers = ['关系类型', '数量', '说明']
    rel_rows = [
        ['RELATES_TO', '4,204', '实体之间的关系（核心知识图谱）'],
        ['MENTIONED_IN', '5,354', '实体被提及在哪些文本块中'],
        ['FROM_DOCUMENT', '179', '文本块来自哪个文档'],
        ['总计', '9,737', '-']
    ]
    add_table_to_doc(doc, '3.2 关系统计', rel_headers, rel_rows)

    # 保存文档
    doc.save(str(docx_path))
    print(f"文档已保存: {docx_path}")

if __name__ == '__main__':
    main()
